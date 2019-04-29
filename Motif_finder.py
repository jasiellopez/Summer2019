import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
import operator
import collections

def main():

    motifs = input("Please input motifs you would like to check for, with a space in between each motif: ")
    motifs_dict1 = {}
    start_site_dict = {}
    motifs_input = motifs.split()
    for motif in motifs_input:
        if "N" in motif:
            motifs = GenerateNComplement(motif)
            for i in motifs:
                motifs_dict1[i] = {}

        else:
            complement = GenerateComplement(motif)
            motifs_dict1[motif] = {}

    files = input("Please input the files you wish to search, with a space between each file path: ")
    files_list = []
    files_list = files.split()
    sequences_dict = {} #dictionary where key is alignment and value is list of alignment
    files_dict = {} #dictionary where the key is the file and vaule is its alignments

    window = int(input("Please input window size you would like to search: "))
    motif_size = int(input("What is the size of the motifs which you want to search for? "))
    newfiles_list = []
    for filename in files_list:
        #/Users/jasiel/Desktop/Molgulid.mesp.alignments.edited_highlighted_motifs.docx
        new_filename = filename + "_highlighted_motifs.docx"
        new_file = docx.Document()
        new_file.save(new_filename)
        newfiles_list.append(new_file)
        file = open(filename, "r")
        sequence_as_list = file.readlines()
        '''
        num_times = 0
        p = new_file.add_paragraph()
        list_styles = new_file.styles
        highlighted = list_styles.add_style('highlighted', WD_STYLE_TYPE.CHARACTER)
        highlighted.base_style = list_styles['Default Paragraph Font']
        highlighted.font.highlight_color = WD_COLOR_INDEX.YELLOW
        string_sequence = ''.join(sequence)
        for motif in motifs_list:
            curr_motif = string_sequence.split(motif)
            for i in range(len(curr_motif) - 1):
                if num_times == 0:
                    p.add_run(curr_motif[i])
                    if len(curr_motif) > 1:
                        motif_run = p.add_run(text=motif, style=highlighted).font
                        motif_run.highlight_color = WD_COLOR_INDEX.YELLOW
        new_file.add_paragraph(p.text)
        new_file.save(new_filename)
        '''
        alignment = ""
        i = 0
        while i < len(sequence_as_list):
            if len(sequence_as_list[i]) > 0 and sequence_as_list[i][0] != '>':
                one_line += sequence_as_list[i].strip("\n")
                if len(sequence_as_list[i]) < line_length:
                    sequences_dict[sequences_name] = one_line
                i += 1
            elif sequence_as_list[i][0] == '>':
                #get sequence length
                sequences_name = sequence_as_list[i]
                split_sequence_name = sequences_name.split()
                #get line_length
                line_length = int(split_sequence_name[1].strip('\n'))
                one_line = ""
                i += 2
        files_dict[filename] = sequences_dict

    for motif in motifs_dict1:
        for filename in files_dict:
            motifs_dict1[motif][filename] = {}
            start_site_dict[filename] = {}
            for sequence in files_dict[filename]:
                motifs_dict1[motif][filename][sequence] = []
                start_site_dict[filename][sequence] = 0

    motifs_found1 = []
    highlighted_sequences = {}
    for filename in files_dict:
        highlighted_sequences[filename] = {}
        for name in files_dict[filename]:
            sequence1 = files_dict[filename][name]
            highlighted_sequences[filename][name] = sequence1
            i = len(sequence1) - 1
            j = 0
            atgFound = False
            while i in range(len(sequence1)):
                if atgFound == True:
                    if (i - 4) >= 0:
                        potential_motif1 = sequence1[i - 4:i]
                        n_occurence = 0
                        for char in potential_motif1:
                            if "N" == char:
                                n_occurence += 1
                        if n_occurence < 2:
                            complement = GenerateComplement(potential_motif1)
                            CompSeq = False
                            inSeq = False
                            if potential_motif1 in motifs_dict1 or complement in motifs_dict1:
                                if potential_motif1 in motifs_dict1:
                                    inSeq = True
                                if complement in motifs_dict1:
                                    CompSeq = True
                                if inSeq == True:
                                    if filename in motifs_dict1[potential_motif1]:
                                        if name in motifs_dict1[potential_motif1][filename]:
                                            motifs_dict1[potential_motif1][filename][name].append(i)

                                if CompSeq == True:
                                    if filename in motifs_dict1[complement]:
                                        if name in motifs_dict1[complement][filename]:
                                            motifs_dict1[complement][filename][name].append(i)
                    i -= 1
                    j -= 1
                else:
                    potential_site = sequence1[(i - 2):(i + 1)]
                    if potential_site == "ATG" or potential_site == "CAT":
                        start_site = i - 2
                        start_site_dict[filename][name] = start_site
                        atgFound = True
                    if atgFound == False:
                        i -= 1
                    elif atgFound == True:
                        i = start_site

    cluster = {}
    window_in_list = {}
    for motif in motifs_dict1:
        cluster[motif] = {}
        window_in_list[motif] = {}
        print("For the motif " + motif + ":")
        for filename in files_dict:
            cluster[motif][filename] = {}
            window_in_list[motif][filename] = {}
            print("\tIn file " + filename + ":")
            for name in files_dict[filename]:
                print("\t\tIn sequence " + name.strip("\n") + ":")
                cluster[motif][filename][name] = {}
                window_in_list[motif][filename][name] = []
                sequence1 = files_dict[filename][name]
                motif_list = motifs_dict1[motif][filename][name]
                motif_index = 0
                cluster[motif][filename][name] = ""
                concatenated_windows = ""
                for index in motifs_dict1[motif][filename][name]:
                    downstream = start_site_dict[filename][name] - index
                    #print("\t\t\t" + str(downstream) + " base pairs upstream of the start site," +
                         # "\n\t\t\tthese are the following motifs within a " + str(window * 2) + " base pair window: ")
                    start = index
                    k = start
                    curr_window = sequence1[start:(index + window + 1)]
                    if motif_index < (len(motif_list) - 1):
                        upstream = (motif_list[motif_index] - motif_list[motif_index + 1]) - 4
                    else:
                        upstream = index - 4
                    not1st_downstream = (motif_list[motif_index - 1] - motif_list[motif_index]) - 4
                    #prevents first window from overlapping start site and also prevents overlap in upstream
                    if upstream >= (2*window) and downstream >= window and motif_index == 0: #done
                        curr_window = sequence1[start - window - 4:start - 4] + sequence1[start:start + window]
                        #print(curr_window)
                        concatenated_windows += curr_window
                        window_in_list[motif][filename][name].append(curr_window)
                        #CheckWindows(concatenated_windows, window, cluster)

                    elif upstream >= (2*window) and downstream < window and motif_index == 0:
                        #done
                        curr_window = sequence1[start - window - 4:start - 4] + sequence1[start:start + downstream]
                        #print(curr_window)
                        concatenated_windows += curr_window
                        window_in_list[motif][filename][name].append(curr_window)
                        #CheckWindows(concatenated_windows, window, cluster)

                    elif upstream < (2*window) and downstream >= window and motif_index == 0:
                        #done
                        if upstream > window:
                            beg_of_upstream_window = 0
                            if motif_index < (len(motif_list) - 1):
                                beg_of_upstream_window = start - 4 - (upstream - window)
                            else:
                                if (start - 4 - window) < 0:
                                    beg_of_upstream_window = 0
                                else:
                                    beg_of_upstream_window = start - window - 4
                            curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:start + window]
                            #CheckWindows(concatenated_windows, window, cluster)
                            #((motif_list[motif_index] - motif_list[motif_index + 1])) gives a positive value since the indexes are stored from the end of the string to the beginning

                        else:
                            if motif_index < (len(motif_list) - 1):
                                beg_of_upstream_window = start - 4 - (upstream - window)
                            else:
                                if (start - 4 - window) < 0:
                                    beg_of_upstream_window = 0
                                else:
                                    beg_of_upstream_window = start - window - 4
                            curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:start + window]

                        #print(curr_window)
                        concatenated_windows += curr_window
                        window_in_list[motif][filename][name].append(curr_window)
                            #CheckWindows(concatenated_windows, window, cluster)

                    elif upstream < (2*window) and downstream < window and motif_index == 0:
                        #done
                        if upstream > window:
                            if motif_index < (len(motif_list) - 1):
                                beg_of_upstream_window = start - 4 - (upstream - window)

                            else:
                                if (start - 4 - window) < 0:
                                    beg_of_upstream_window = 0
                                else:
                                    beg_of_upstream_window = start - window - 4

                            curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:start + downstream]
                            #CheckWindows(concatenated_windows, window, cluster)
                        else:
                            if motif_index < (len(motif_list) - 1):
                                curr_window = sequence1[start:start + downstream]
                            else:
                                if (start - 4 - window) < 0:
                                    beg_of_upstream_window = 0
                                else:
                                    beg_of_upstream_window = start - window - 4
                                curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:start + downstream]

                        concatenated_windows += curr_window
                        #print(curr_window)
                        window_in_list[motif][filename][name].append(curr_window)
                        #CheckWindows(concatenated_windows, window, cluster)

                    if motif_index > 0:
                        #if both windows are large enough to allow for normal window
                        if upstream >= (2*window) and not1st_downstream >= (2*window):
                            #done
                            curr_window = sequence1[start - window - 4:start - 4] + sequence1[start:start + window]
                            concatenated_windows += curr_window
                            #print(curr_window)
                            window_in_list[motif][filename][name].append(curr_window)
                            #CheckWindows(concatenated_windows, window, cluster)

                        #if window upstream is large enough to allow for a normal window to be taken but window downstream will lead to overlap unless adjusted
                        elif upstream >= (2*window) and not1st_downstream < (2*window):
                        #done
                            if not1st_downstream >= window:
                                curr_window = sequence1[start - window:start - 4] + sequence1[start:start + window]
                                #CheckWindows(concatenated_windows, window, cluster)
                                #((motif_list[motif_index] - motif_list[motif_index - 1])) gives a positive value since the indexes are stored from the end of the string to the beginning

                            else:
                                if not1st_downstream > 0:
                                    curr_window = sequence1[start - window:start - 4] + sequence1[start:start + not1st_downstream]
                                else:
                                    curr_window = sequence1[start - window:start - 4]

                            concatenated_windows += curr_window
                            #print(curr_window)
                            window_in_list[motif][filename][name].append(curr_window)
                            #CheckWindows(concatenated_windows, window, cluster)

                        #if both windows up and downstream will lead to overlap unless adjusted
                        elif upstream < (2*window) and not1st_downstream < (2*window):
                            #done
                            if not1st_downstream >= window:
                                end_of_downstream_window = start + window

                            else:
                                if not1st_downstream > 0:
                                    end_of_downstream_window = start + not1st_downstream
                                else:
                                    end_of_downstream_window = start

                            if upstream > window:
                                #gets window upstream of size that allows for the next index to get downstream window of size of window variable without overlapping
                                if motif_index < (len(motif_list) - 1):
                                    beg_of_upstream_window = start - 4 - (upstream - window)
                                else:
                                    if (start - 4 - window) < 0:
                                        beg_of_upstream_window = 0
                                    else:
                                        beg_of_upstream_window = start - window

                            else:
                                if motif_index < (len(motif_list) - 1):
                                    beg_of_upstream_window = start - 4
                                else:
                                    if (start - 4 - window) < 0:
                                        beg_of_upstream_window = 0
                                    else:
                                        beg_of_upstream_window = start - window - 4

                            curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:end_of_downstream_window]
                            concatenated_windows += curr_window
                            #print(curr_window)
                            window_in_list[motif][filename][name].append(curr_window)
                            #CheckWindows(concatenated_windows, window, cluster)

                        #if window downstream is large enough to allow for a normal window to be taken but window upstream will lead to overlap unless adjusted
                        elif upstream < (2*window) and not1st_downstream >= (2*window):
                            #done
                            if upstream > window:
                                if motif_index < (len(motif_list) - 1):
                                    beg_of_upstream_window = start - 4 - (upstream - window)
                                else:
                                    if (start - 4 - window) < 0:
                                        beg_of_upstream_window = 0
                                    else:
                                        beg_of_upstream_window = start - window - 4

                                curr_window = sequence1[beg_of_upstream_window:start - 4] + sequence1[start:start + window]
                                #CheckWindows(concatenated_windows, window, cluster)
                                #((motif_list[motif_index] - motif_list[motif_index - 1])) gives a positive value since the indexes are stored from the end of the string to the beginning

                            else:
                                if motif_index < (len(motif_list) - 1):
                                    curr_window = sequence1[start:start + window]
                                else:
                                    if (start - 4 - window) < 0:
                                        curr_window = sequence1[0:start - 4] + sequence1[start:start + window]
                                    else:
                                        curr_window = sequence1[start - window:start - 4] + sequence1[start:start + window]
                            concatenated_windows += curr_window
                            #print(curr_window)
                            window_in_list[motif][filename][name].append(curr_window)
                                #CheckWindows(concatenated_windows, window, cluster)

                    motif_index += 1

                    #for co_motif in cluster[motif][filename][name][index]:
                        #print("\t\t\t\t" + co_motif + " following distances (negative represent upstream):")
                        #print("\t\t\t\t\t" + str(cluster[motif][filename][name][index][co_motif]))
                cluster[motif][filename][name] = concatenated_windows
                if len(window_in_list[motif][filename][name]) > 0:
                    co_motifs = FindRandMotifs(concatenated_windows, motif_size)
                    #print(filename)
                    #print(name)
                    co_motifs1 = FindRandMotifs(window_in_list[motif][filename][name], motif_size)
                    print("List ")
                    for key in co_motifs1:
                        print(str(key) + ": " + str(co_motifs1[key]))
                    print("\nString ")
                    for key in co_motifs:
                        print(str(key) + ": " + str(co_motifs[key]))
'''
    for filename in files_dict:
        for name in files_dict[filename]:
            sequences = FindSeqWithMotifs("GATA", 2, "GGAA", 1, files_dict[filename][name], 30, start_site_dict[filename][name])
            print(sequences)
'''
#change function names to proper format "def word_word_word"

#gives index of co-motif occurences around a certain window of a motif
def CheckWindows(concatenated_window, window, cluster):
    while k in range((index - window), (index + window + 1)):
        curr_potential_motif1 = sequence1[k:k + motif_size]
        complement = GenerateComplement(curr_potential_motif1)
        if curr_potential_motif1 in cluster[motif][filename][name][index]:
            cluster[motif][filename][name][index][curr_potential_motif1].append((index - k))
        elif complement in cluster[motif][filename][name][index]:
            cluster[motif][filename][name][index][complement].append((index - k))
        else:
            cluster[motif][filename][name][index][curr_potential_motif1] = [(index - k)]
        k += 1

def PrintHighlighted(motifs_dict1):

    for motif in motifs_dict1:
        for filename in files_dict:
            highlighted_sequences[filename] = {}
            for name in files_dict[filename]:
                print("\033[1;33m"+motif+"\033[3;37m")
                sequence = ""
                sequence = str(files_dict[filename][name])
                highlighted_motif = "\033[1;33m" + motif + "\033[3;37m"
                sequence = sequence.replace(motif, highlighted_motif)
                complement = GenerateComplement(motif)
                highlighted_complement = "\033[1;33m" + complement + "\033[3;37m"
                sequence = sequence.replace(complement, highlighted_complement)
                files_dict[filename][name] = sequence
                print(name)
                print(files_dict[filename][name])

    return sequence

def GenerateComplement(sequence):

    complement = ""
    i = len(sequence) - 1

    while i >= 0:
        if sequence[i] == "A":
            complement += "T"
        elif sequence[i] == "G":
            complement += "C"
        elif sequence[i] == "C":
            complement += "G"
        elif sequence[i] == "T":
            complement += "A"
        elif sequence[i] == "N":
            complement += "N"

        i -= 1

    return complement

def GenerateNComplement(sequence):

    complements = ["", "", "", ""]
    i = len(sequence) - 1
    n_occurence = 0
    for char in sequence:
        if "N" == char:
            n_occurence += 1
    if n_occurence < 2:
        while i < len(sequence) and i >= 0:
            if sequence[i] == "A":
                complements[0] += "T"
                complements[1] += "T"
                complements[2] += "T"
                complements[3] += "T"

            elif sequence[i] == "G":
                complements[0] += "C"
                complements[1] += "C"
                complements[2] += "C"
                complements[3] += "C"

            elif sequence[i] == "C":
                complements[0] += "G"
                complements[1] += "G"
                complements[2] += "G"
                complements[3] += "G"

            elif sequence[i] == "T":
                complements[0] += "A"
                complements[1] += "A"
                complements[2] += "A"
                complements[3] += "A"

            else:
                complements[0] += "A"
                complements[1] += "T"
                complements[2] += "C"
                complements[3] += "G"

            i -= 1

    return complements



def FindRandMotifs(sequence, motif_size):

    most_common_motifs = {}
    motifs_found1 = {}
    motif_indexes = {}

    if type(sequence) is str:
        i = 0
        print(sequence)
        while i <= (len(sequence) - motif_size) and len(sequence) > 3:
            curr_potential_motif1 = sequence[i:i + motif_size]
            n_occurence = 0
            for char in curr_potential_motif1:
                if "N" == char:
                    n_occurence += 1
            if n_occurence < 2:
                complement = GenerateComplement(curr_potential_motif1)
                if curr_potential_motif1 in motifs_found1:
                    motifs_found1[curr_potential_motif1] += 1
                    motif_indexes[curr_potential_motif1].append(i)
                elif complement in motifs_found1:
                    motifs_found1[complement] += 1
                    motif_indexes[complement].append(i)
                else:
                    motifs_found1[curr_potential_motif1] = 1
                    motif_indexes[curr_potential_motif1] = [i]
            i += 1

        for j in range(10):
            key = max(motifs_found1.items(), key=operator.itemgetter(1))[0]
            motifs_found1.pop(key)
            most_common_motifs[key] = motif_indexes[key]

    if type(sequence) is list:
        index = 0
        for window in sequence:
            declared = False
            if len(window) > 3:
                sequence1 = window
                #print(window)
                i = 0
                #print(len(sequence1) - motif_size)
                while i <= (len(sequence1) - motif_size):
                    curr_potential_motif1 = sequence1[i:i + motif_size]
                    #print(curr_potential_motif1)
                    n_occurence = 0
                    for char in curr_potential_motif1:
                        if "N" == char:
                            n_occurence += 1
                    if n_occurence < 2:
                        complement = GenerateComplement(curr_potential_motif1)
                        if curr_potential_motif1 in motifs_found1:
                            motif_indexes[curr_potential_motif1].append(index)
                            motifs_found1[curr_potential_motif1] += 1
                        elif complement in motifs_found1:
                            motif_indexes[complement].append(index)
                            motifs_found1[complement] += 1
                        else:
                            motif_indexes[curr_potential_motif1] = [index]
                            motifs_found1[curr_potential_motif1] = 1

                    i += 1
                    index += 1
                index += motif_size - 1

        for j in range(10):
            key = max(motifs_found1.items(), key=operator.itemgetter(1))[0]
            motifs_found1.pop(key)
            most_common_motifs[key] = motif_indexes[key]

    return most_common_motifs


def FindSeqWithMotifs(motif1, num_motif1, motif2, num_motif2, sequence, window_to_search, start_site, motif_size):

    #list of sequences that fit the criteria
    list_sequences = []
    print(sequence)
    if start_site > 0:
        i = 0
        while i in range((start_site - window_to_search)):
            window = sequence[i:i + window_to_search]
            num_occurrences1 = 0
            num_occurrences2 = 0
            for j in range(len(window) - 3):
                potential_occurence = window[j:j + motif_size]
                if potential_occurence == motif1 or GenerateComplement(potential_occurence) == motif1:
                    num_occurrences1 += 1
                elif potential_occurence == motif2 or GenerateComplement(potential_occurence) == motif2:
                    num_occurrences2 += 1
            if num_occurrences1 == num_motif1 and num_occurrences2 == num_motif2:
                list_sequences.append(sequence[i:i + window_to_search])
            i += 1
    return list_sequences

main()

'''
THIS IS SUPPOSED TO HIGHLIGHT ALL THE MOTIFS BUT IT ISN'T FIGURE IT OUT
    for filename in files_dict:
        for sequence in files_dict[filename]:
            highlighted_sequence = ""
            for motif in motifs_dict1:
                unhighlighted = files_dict[filename][sequence].split(motif)
                highlighted1 = "\033[1;33m"+motif+"\033[3;37m"
                for j in range(len(unhighlighted)):
                    if j != (len(highlighted1) - 1):
                        highlighted_sequence += unhighlighted[j] + highlighted1
                    else:
                        highlighted_sequence += unhighlighted[j]
                highlighted_sequences[filename][name] = highlighted_sequence
                print(files_dict[filename][sequence])
'''



    #print("\n")
    #print(cluster)
    #concatenate all windows and search for enrichment in them
    #clustering: getting a sequence with n amount of one motif and m amount of other co-motif within x amount of base pairs
    #secomod script
    #use neural genes to test algorithm (test for gata and ets site)
    #look up nucleosome finding script

'''paragraphs = []
    #TODO: DO OTHER STUFF
    #inserts a -1 when a new alignment sequence is starteds
    for motif in motifs_found1:
        motifs_dict1[motif].append(-1)
    for motif in motifs_found2:
        motifs_dict2[motif].append(-1)
    motifs_found1 = []
    motifs_found2 = []
    for key in motifs_dict1:

correction = 0
for i in range(0, (len(sequence1)//60)):
    print((len(sequence1)//60))
    print("IN IF STATEMENT")
    print(i)
    print(len(sequences_dict[key][0]))
    corrected_i = (i*60) + correction
    print(corrected_i)
    pre_new_line = len(sequences_dict[key][0])
    print(sequences_dict[key][0][:50])
    print(sequences_dict[key][0][:5motif_size])
    print(sequences_dict[key][0][:58])
    print(sequences_dict[key][0][:62])
    sequences_dict[key][0] = sequences_dict[key][0][:corrected_i] + "\n" + sequences_dict[key][0][corrected_i:]
    print(len(sequences_dict[key][0]))
    print(sequences_dict[key][0][:corrected_i])
'''
'''print(sequences_dict[key][1])
    sequences_dict[key][1] = sequences_dict[key][1][:corrected_i] + "\n" + sequences_dict[key][1][corrected_i:]
    print("\n")
    print(sequences_dict[key][1])

'''
        #correction = len(sequences_dict[key][0]) - pre_new_line

#add two lines at a time in a run (current line and alignment)

#find distance between 2 motifs
#identify if there is a section of N bp where a certain amount of motifs appear



#consult andrew
#find overrpresentation close to an identified ETS site

#ask joe whether we shouldve been done with specs we ran












'''import docx

def main():

    motifs = input("Please input motifs you would like to check for, with a space in between each motif")
    motifs_dict = {}
    motifs_list = motifs.split()
    for motif in motifs_list:
        motifs_dict[motif] = []

    files = input("Please input the files you wish to search, with a space between each file path")
    files_list = []
    files_list = files.split()
    sequences_dict = {}

    for filename in files_list:
        file = open(filename, "r")
        sequence = file.readlines()
        one_line = ""
        alignment = ""
        i = 0
        while i < len(sequence):
            if len(sequence[i]) > 0 and sequence[i][0] != '>':
                one_line += sequence[i].strip("\n")
                alignment += sequence[i + 1].strip("\n")
                if len(sequence[i]) < 60:
                    sequences_dict[sequences_name] = [one_line, alignment]
                i += 3
            if i < len(sequence):
                if sequence[i][0] == '>':
                    sequences_name = sequence[i]
                    i += 2
    print(sequences_dict)

    newfiles_list = []
    for i in len(sequences_dict):
        for
        #/Users/jasiel/Downloads/Molgulid.mesp.alignments.edited_highlighted_motifs.docx
        new_filename = files_list[i] + "_highlighted_motifs.docx"
        new_file = docx.Document(new_filename)
        remainder = 0
        curr_line = 0
        #denotes whether previous line was divisible by motif_size
        previous_line = False
        for line in file:
            if ((len(line) + remainder) % motif_size != 0):
                remainder = len(line) % motif_size
                for j in range(len(line)):
                    if j <= len(line) - remainder - 1:

                        potential_motif = line[j:j + motif_size]
                        #potential_motif =
                        if potential_motif in motifs_dict:
                            found_motif_index = j;
                            motifs_dict[potential_motif].append(found_motif_index)
                            #highlighttext
                            #increment occurence amount
                    if j > len(line) - remainder - 1:
                        sequence_remainder = line[j:]
                previous_line = True
            else:
                for j in range(len(line)):
                    potential_motif = line[j:j + motif_size]
                    if potential_motif in motifs_dict:
                        found_motif_index = j;
                        motifs_dict[potential_motif].append(found_motif_index)
                        #highlighttext
                        #increment occurence amount
                previous_line = False
            curr_line += 1

    print (motifs_dict)
#find distance between 2 motifs
'''




'''
>>CI-OTX 60

CGTTATCTCTAACGGAAGTTTTCGAAAAGGAAATTGTTCAATATCTAAGATAGGAATG

>>Co-Otx 60

CCGGTAATTACTCCATCAACAATAATTCAAGTTCATCTCAATTGAGGTAAGGAAAAACT
TCAATTTCTTAATTGAACAAATTTATTTAAAATCTAAATTGTTTATCAAGACGAGATGT
TTACTCTTTCGGTACACATTTATATTATAAACAAATTCCCCTTATGTTATGGGAACTAA
TATTCGAAGGCGTCTGGTCATTTATAGCAAGTAAAAAATACGTAAATAAATAAACAAAA
TACGAAGCCCGGAATAATAACATGCGATGAATGTATTCGGGTTTTGCTATTGTGTTACG
TAAAGAGCGCATTTTTTGTGTCAATAAATATTCTGCATTAGAATAAACGACAAAGCTAT
TGCACCATGGGAAAACTGCGTAGTAGGATTTCTTTTGGACAACGTAATAAGTTTGCGCA
ATTTTTCTTACAATTTGCAGCTCTAAAATATTTATAAAGACGGATATGAATTTAAAAAT
GATTTAAATTTATGAATTAATATATGGAATTAAGTTATTTCTAATAGATCATTTTCGGT
TGATAATATAATTCATATGAATTATTATATTAAATTTAGTTATAGTTAATAGATCATTT
TAGGATTTATAAATTAGTCTTAGCCATAAAAATATTCAGAAAATAATATTTGACCGCTG
GTCGCTCACGCATCACGTGATCACCATCCTGTTTAAATTATTGGACAAGATTCGTTTAA
GTTTACATAGGGTTCCTTGGTAGAGGATTAACTTTATTAATCCCTATAATTACCTTTGA
TTTCATTCGAATGTCGACCATTGTTCGGAAGCCAAGCAGAATGTTAACCAACGAAATCA
AAACCGTATTTTGACGCGTGTACTAATGCATAATTTGTCCATTGTGCTGCTTTATTTGC
TTTGGAGCCCAGAAAACGTGCACCGGTTTTTGCAGGCAGTAATTAGTCGCCGAAACAAA
AGTTCTAAAGTTCGTTTTCCGCACTCTTATAAACTACCTTGCAGCGCGCACATGTATGT
AATGAGGACAGCCGGTTATAACAATGCATCGGCAAGGTGCCGTATTTGAATAGGCCTTC
TCAGAGAGTGTGTGAAACTTTAAGTAATAGTTAATAAGAGTAAATCGTTAATATACGTA
TAAGCCGATCAGAAGGATATAACTTATTAAAATTATATAACTCTTTTTGTTTCAGGTTT
TAAGTTGAGTTTTTGTCAATTCACAAACGGAAATTATGTCTTATTTGAAGCCACCTCAT
TATGCCATG

>>Moccul.mesp.enhancer 60

AATCGGTGAAACCATATATTTCTATATATGTTGGGATTCGGTGGGATGAAATTATGTTA
TTTAGAACCGACTTACATTTGTGCATCTCTGACTGACTTGCACGACAAATATCGTTGTA
CCATCGGTTGGTTGATGGACGCCATCATCACCCAAATGATTGCACCACCGATCCCAAGT
GACCATGACGTGTGTTTGATGTATGGATCCGTGTCCCAACTGAAACACGAAATTGTTAA
AACGTAATATTAAACGTTGCAAATGGACAAGAGACTGAAAGTAGATGACAGCTTAAACT
AAATAATTTGTCAAGTGGCCGTTTGTTATAGTTTTGTATGTTTTGGTTAAATATTTCAC
TTCTTACTTGAAGAAATTGTCTCTTTTCCCAATTCTCATGGCCTCCAGAACTTGATCAA
AACCCCCAACAGTTGCACTTGTCTTTGCAATAACAAGAATCATTCCAAGGACCATTACA
CCAGCCTGAACAACATCTGTCCATATGACTCCTTTGATGCCGCCCTGTGAAGAAGAGGG
TTGAAACTATTGCATCAAACACATTAACAGAGAAATTCAACATTATACATCGTACCAAA
AAGTGTGTTGTGACAAAGATTTTAAAATTTTACAACTATAATACCGATCCGGAATGTTA
TTTCAGAAGAACCGAGGTACTGTCGGTAAGTGCTTGAAGGTCAAGAAATTGGCGAAGTA
TTGCACAGGTAGCAAAAGATGCAACATCTCGATTCTCGACGTAAAGTTGCAGACAATTA
ATTAATTAAACACTTCCGACATGATTGATCATCGCCGTTATCAGGAAAACTGATAACAT
TATCGCAAGTTTTATCGAGAACTGATTGCAAGGCGATAAGACTGGATTCCCTTCAACCA
AGATTGAAATAGCAGGAGGCAAAATGGGGTGAATTATGTCTGGACAAATGACAACTTTA
TGTACTATAAATCACGCAATTGTTTGTCATAAGCACACATAGTCAACGATAAAGTTATT
AATTCATTAAATTTAATCGTTTATATTATGAACAAACAAACGATTGTCATGAGAAGTTA
TTGCGAAGAAACAATCCAGTCTGTTCGATTGCAATGTCCAAATCAACAAACGAGCGTTC
CCGATGTATTTACTGACCGAGTTCCAAACAAACAAACCTTTATTGCGTATCCTGGGTCT
ACTGCTTCACATTTCCCTGC

>>Moocul.mesp.enhancer 60

NNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNTGTGTCGGATCAAGATCTCA
TCTCTACCCCAAACTCAAACCCCACCTACACTTGAGGTATTTGTCACAACGACATAACA
ACAAGAAATGGACGGATTGAGTGAAATTTCTAATGAATTTGTAATCAACTTATTCGTTT
TATTGACATCATTTGGCAACTGATGGTATATGGATCCATTATAGCACATCATACCATTA
CTTTATAATGGAATCCATACCATTAATGTTACTTGCGATGCGCACAGTGTAGCATATAC
ATAAGAATGGTTTCAGGCTATTCGACTATTGTACATACATTTGTGCGTCTCTCATTGAC
TTGCATGACAAATATCGTTGTACCATCGGTTGATTAATAGCTGCCATCATCACCCAGAT
GATGGCGCCACCGATCCCCAGTGACCAAGATGTGTGTTTGATGTATGGATCCGTATCCC
AACTGAAACATGAAATTGTGAAGCGTACTCAGTGCTGGGCAATAAACTGAAACGATGTG
AAAACATCCACTGCAAAATTTCAGTACAATTCGCATGGTTCTGATTAAGATAATGGAAT
ATTACAATTTCAGTAATAGAGTTGATAACGGAAAATAAGTATACTTATGGGTTTATACT
TATTTTATTATAAAGGTTGTTAAAAAAGTTAAAAAATAGTTGTGGATATCGGTCCAAAA
ACTTAAAATAACCCATCAAATAGCTTTCAATTCGTACTTGAAGAAGTTGTCTCTTTTTC
CAACCCTCATCGCCTCCATCACTTGATCAAAGCCTCCAACAATTGCACTTGTCTTCGCA
ATCACAAGAATCATTCCGAGTATCATGACGCCAGCCTGCACGACATCCGTCCATATGAC
TCCTTTGATGCCACCCTGTTGAAAGTGAAAGGTTATGTGATCGGTACGCAGATAAAGTA
TGGGAAATAAGACAGTCTTACATGACCGGTCACGGTCATGTAACCTAAACCATGACTTA
CACAGATCAAGATCTATTACAACATGCTATATACGGTTATAATATACATTGTGTTTCAA
TTGTCAATTAAGTACCACGAACCATCCATCACATTAAAGTTAATACTTTATAACAGAAG
CCAGCATTAATTCTGTCAAAAATATTTGCAATATTCACTGGTGGTTTCGATTTTAAGAT
ATCACAAATAATGTTTTCCAAAAATTTCTGTCAATGAAACCGAGATGTTGGATTACGTC
GGTAATTAGTAGCAGAAGTCGATGCAATATATCTCATCACATATGTTTACAGACAAATT
AATTAAACCCTTCCGACATGATTGGTCATCACCGCTATCACTAAAGCTGATAACATTAT
CGCAAGTTTTATCGAAAACTGATTGCAAGGCGATAAGACTGGATTGTCTTCGACGAAGC
TTGAAATAGCAGGAGGCAAAACGGAGTGAATTATGTCTGGACAAATACTGACTCTCTGT
TCTATAAATTCCTTCATTGTTTGTCATAAGCACACATTGTCAACGATAAAGTTATTAAT
TCATTAAATTTAATCGTTTATATTATGAACAAACAAACGATTGTCATGAGAAGTTATTA
TGAAGAAACAATCCAGTCAGTATCGTCACAATGTCCGAATCAACAAGCGGGCGTTCCCG
GTGTATTTACCACCCAAGTTCCAAACAAACAAACGTTTATTGCGTATC

>>Moocci.mesp.enhancer 60

GTGTTTAGCGATCACGGTTCAAGTTTAGGCTACATCGCTTTTTCCATCGGTTTGCCTAA
ACTTGCATGCTGTATACAGTATTAAGAAGGTCTTGTTTAACAATGTTCCTTAAATGTTA
TTGACATACAAAGCAAACTTTGTGTTAAAAGCCTACTTTAATGTGAAGACCAACATAAA
AAAATTATCGGATAGTATAACTTACTCAAAAAAGTTAAGTCTTTGTCCTCTTTGCAAAG
CTGCAGTAACTTTATCGAATCCCCCAAGATACACGGATGTTCTGATGATAACTAGAAGC
ATCCCCACTATCATTACTCCAGCTTGTACAACGTCTGTCCATATCACACCTTTTAAACC
ACCCTGGAAATTAAATCATTATTATTATTAAAGATAAAAACCATTTATAAGCCTATTTA
CAATACAGTTCCAAATTTAATCTGCCAATTCGATTGAACGTTTTTCAATGAAGTAAAAT
ATCTCTCACAATCTTGTATTGAGCCTTTAAAATAGCCTTTTAAATGAAAATCTATATCT
GACAAAATAGGCTTTGACGTGATTTGAACGCCGTACCTCATGCTCGTATAAATACCTAA
TTTGTTGCGGAATTTCTACTTCAATATAAAAAGTATGCAAATACATACATAATAGTTAT
TGAAAAGAATTTAAATCATCAAATTTTATGTCTGTTGTAGCCTAAATAGTGTTTTATTT
AAGTAAATCTACACTGCGCTTCAGTTCAAACCAACCGTGCCTTGAACTACGGGCAGTGC
GGTTGCTATAGACGTTATGCTTACGTAAACTAGTGCAAAGAATTTGTTGGAAAGTGTTA
ATTGATATCCGACCACCGACTGTGTTATCTTATCGTTTTGATAGATAAGAGCTGATAAC
ACTGATAATAGACTAATTAACTATAGCAGTTAATTTTTATTAAAATTTTCACTTTTCGG
ACATGAAAATTTGTCTAAAACCCAAACTTCCAAACTTCATCTATATAAGGAAAGAATTT
TATTGAAATGGCATACTCAATTCATCAGTATATATCATCTATATCTATATACGGTGGGT
AATACGTCGTAATTGATATCGGACAAAGTAAAAACTGCATATTGAAGCGTTAGAAGAAA
GTTGTTCAATTCAACATTTTAACAAAATG

>>Cirobu.mesp.enhancer 60

TTTTACATTTGAAATGTGATTAATTACGAAAATCCAGCGAATAGAATTGTCACAACAAG
TCATTAGCGACGGATATTTCGCCTTTGAAACTTAAAGGCGATAATGACTTTGCCCGTTT
CATGCGGCGATAAACGAACTAATTAGACACCTCCTACAGATATAATGGTAATTCAGAAT
CGTGTGGTTATGTAATTCACAAAAACATTTTAACAAAACAGGTTGATTTGAAACTTGTA
TTATGGAACTCAGTATGCATCAGGTAAGACGACGTGTACAATTTTTTCCAAAATAACGA
ATATA
''
